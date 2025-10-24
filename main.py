import os
import requests
import json
import time
import re
import queue
import threading
import uuid
import shutil
from datetime import datetime
from io import BytesIO
from flask import Flask, render_template_string, request, jsonify, send_file
from markdown import markdown
# NOTE: This application requires the 'python-docx' and 'beautifulsoup4' libraries.
# Install with: pip install python-docx beautifulsoup4
from docx import Document
from bs4 import BeautifulSoup

# --- LLM API Configuration ---
# 请根据你的环境配置这些设置
VLLM_SERVER_HOST = "ip"
VLLM_SERVER_PORT = "port"
OPENAI_API_ENDPOINT_PATH = "/v1/chat/completions"
TARGET_API_URL = f"http://{VLLM_SERVER_HOST}:{VLLM_SERVER_PORT}{OPENAI_API_ENDPOINT_PATH}"
SERVED_MODEL_IDENTIFIER = 'models'
API_HEADERS = {
    "Content-Type": "application/json",
    # "Authorization": "Bearer YOUR_API_KEY" # 如果需要，请添加你的API密钥
}

# --- 所有生成作品的基础目录 ---
WORKS_LIBRARY_PATH = "works_library"
WORLDVIEW_DIR_NAME = "_worldview"
WORLDVIEW_FILE_NAME = "worldview.md"
MIND_LIBRARY_DIR_NAME = "_mind_library"
MIND_LIBRARY_FILE_NAME = "mind_library.json"
TASK_LIBRARY_FILE_NAME = "task.json"
TOKEN_STATS_FILE = "token_stats.json"
TASKS_STATE_FILE = "tasks_state.json"
NEW_TASKS_STATE_FILE = "new_tasks_state.json" # New constant added
OLD_TASKS_STATE_FILE = "old_tasks_state.json" # New constant added

OUTLINE_PROMPT_TEMPLATES = {
    "course": {
        "default": """
        请你扮演一位资深的课程设计专家。
        我需要一个关于 "{topic}" 的课程大纲。
        请遵循以下结构和要求：
        1.  在第一行，请以 "标题：[这里是课程标题]" 的格式，为这个课程给出一个简洁、明确的标题。
        2.  从第二行开始，生成详细的课程大纲。
        3.  大纲应包含 {chapters} 个主要章节。
        4.  每个章节下应包含 {sections} 个小节。
        5.  每个小节都需要有一个明确、具体的标题。
        请以Markdown格式输出完整的大纲，章节使用'###'，小节使用'1.'、'2.'等数字列表。
        """
    },
    "novel": {
        "default": """
        请你扮演一位经验丰富的小说编辑。
        我正在构思一部关于 "{topic}" 的小说，请帮我生成一份故事大纲。
        请遵循以下结构和要求：
        1.  在第一行，请以 "书名：[这里是小说标题]" 的格式，为这部小说想一个引人入胜的标题。
        2.  从第二行开始，生成详细的故事大纲。
        3.  小说分为 {chapters} 个主要章节。
        4.  每个章节标题应能概括该章节的核心事件或情感转折。
        请以Markdown格式输出完整的大纲，章节使用'###'。
        """
    }
}

CONTENT_PROMPT_TEMPLATES = {
    "course": """
    请你扮演一位资深的课程研发专家和讲师。
    你的任务是为 "{topic}" 这个课程下的一个小节撰写一份详细、专业且通俗易懂的教学讲义。

    **小节标题**: "{section_title}"

    **内容要求**:
    1.  **格式**: 严格使用 Markdown 格式。
    2.  **结构**: 内容应包含引言、核心概念讲解、关键技术点、代码或伪代码示例（如果适用）、实际案例分析、以及小结。
    3.  **深度**: 内容需要有足够的深度，覆盖该知识点的核心原理和关键细节。
    4.  **专业性**: 术语使用准确，讲解系统化。
    5.  **篇幅**: 请确保内容详实，提供丰富的示例和深入的解释。

    请开始为 "{section_title}" 这一节撰写讲义内容。
    """,
    "novel": """
    请你扮演一位才华横溢、文笔细腻的小说家。
    你的任务是完全按照给定的“本章剧本”，并融入丰富的文学技巧，创作出小说《{topic}》中指定章节的完整内容。

    **核心任务**: 续写标题为“{section_title}”的章节。

    **背景信息 (供参考)**:
    1.  **宏大世界观 (摘要)**:
        ```
        {worldview_context}
        ```
    2.  **上一章节回顾 (精简摘要)**:
        ```
        {previous_chapter_context}
        ```

    **创作指令 (必须严格遵守)**:
    1.  **本章剧本 (来自思维库)**:
        ```
        {chapter_plan}
        ```
    2.  **绝对遵循剧本**: 你的创作必须**严格且完全**基于“本章剧本”进行。剧本中提到的剧情点、人物互动、场景和要引入的新信息都必须在正文中得到体现。**不要偏离剧本，不要即兴发挥**。
    
    **文笔与技巧要求 (非常重要)**:
    1.  **展示而非告知 (Show, Don't Tell)**: 不要直接说“他很害怕”，而是通过描写他“手心出汗、呼吸急促、眼神躲闪”来展现他的害怕。用动作、场景、感官细节来传达情感和信息。
    2.  **丰富感官描写**: 充分描写角色所看到的、听到的、闻到的、尝到的和感觉到的东西，让读者身临其境。
    3.  **避免重复**: 尤其是对于核心设定（如主角的能力），请使用多样化的词汇和比喻进行描述，避免在每一章都使用完全相同的形容词。
    4.  **个性化对话**: 让每个角色的说话方式都符合其在世界观中的设定。一个神秘的导师和一个鲁莽的战士，他们的用词、语气和节奏应该完全不同。对话不仅要推动情节，更要塑造人物。
    5.  **强化内心独白**: 深入主角的内心世界，通过他的思考、挣扎和自我怀疑，来展现人物的成长和弧光，让角色更加立体。

    **最终要求**:
    - 直接创作小说正文，专注于场景、动作、对话和心理描写。不要在正文中包含任何元注释、分析或标题。
    - 保持生动的文笔，并确保本章内容详尽，字数在 {word_count} 字左右。

    现在，请严格依据“本章剧本”和上述所有技巧要求，开始创作《{topic}》中标题为“{section_title}”的章节。
    """,
    "generate_chapter_script": """
    请你扮演一位心思缜密的小说剧情规划师。
    你的任务是基于已有的宏大世界观和刚刚结束的章节，为即将开始的新章节制定一个详细的执行剧本。

    **1. 宏大世界观 (世界规则与核心设定)**:
    ```
    {worldview}
    ```

    **2. 上一章剧情回顾 (精简摘要)**:
    ```
    {previous_chapter_summary}
    ```

    **3. 本章标题**: "{section_title}"

    **你的任务是为 "{section_title}" 这一章，规划出清晰、可执行的写作蓝图。请严格按照以下JSON格式输出**:
    ```json
    {{
      "chapter_title": "{section_title}",
      "status": "pending",
      "summary": "用2-3句话概括本章的核心故事脉络。",
      "key_events": [
        "列出本章必须发生的第1个关键事件。",
        "列出本章必须发生的第2个关键事件。",
        "列出本章必须发生的第3个关键事件。"
      ],
      "登場人物": ["列出本章的主要出场人物"]
    }}
    ```
    请确保输出是**一个完整且格式正确**的JSON对象，不要包含任何额外的解释或Markdown标记。
    """,
    "summarizer": """
    请将以下小说章节内容浓缩成一段150字以内的精简摘要，用于帮助AI理解上下文。摘要需要清晰地概括出以下几点：
    1.  本章的核心事件是什么？
    2.  主要人物在本章结束时的状态（物理、情感）是怎样的？
    3.  是否留下了悬念或为下一章铺垫了哪些线索？

    **章节内容**:
    ---
    {text_to_summarize}
    ---
    
    请输出精简摘要。
    """,
    "initial_worldview": """
    请你扮演一位世界级的小说策划师和设定师。
    你的任务是根据提供的小说主题和高级大纲，构建一个全面且详细的“世界观圣经” (Worldview Bible)。这份文档是后续所有章节创作的基石，必须确保设定的严谨性和一致性。

    **小说主题**: "{topic}"

    **故事大纲**:
    ```
    {outline}
    ```

    **请在世界观圣经中包含以下核心部分，并使用Markdown格式化**:

    ### 核心概念
    用一到两句话总结整个故事的核心冲突和主题。

    ### 主要人物
    为每位主要人物（至少2-3位）建立档案，每个档案使用 '---' 分隔。档案包括：
    - **姓名**:
    - **身份/背景**:
    - **外貌特征**:
    - **性格特点**: (例如：勇敢但鲁莽，聪明但多疑)
    - **核心目标/动机**: (他们最想要什么？)
    - **主要冲突/困境**: (什么在阻碍他们？)

    ### 世界设定
    - **时代与地点**: (故事发生在何时何地？是现代都市，还是架空古代，或是未来太空？)
    - **关键地点描述**: (描述几个故事中会反复出现的关键场景，例如：主角的家、神秘的森林、繁华的都市中心等。)
    - **社会规则/文化背景**: (这个世界有什么独特的法律、习俗、信仰或技术水平？)
    - **氛围与基调**: (故事是光明的、黑暗的、悬疑的，还是幽默的？)

    ### 关键情节节点
    根据大纲，列出几个将推动故事发展的关键转折点或核心事件。

    请确保内容详尽、逻辑自洽，为AI后续创作提供清晰、无歧义的指导。
    """,
    "extract_worldview_updates": """
    请你扮演一位严谨的小说档案管理员。
    你的任务是只从“最新创作的章节内容”中，提取出所有**新增的或发生显著变化**的世界观信息。

    **1. 已有的世界观设定 (仅供参考，不要重复输出已知信息)**:
    ```
    {existing_worldview}
    ```

    **2. 最新创作的章节内容**:
    ```
    {new_content}
    ```

    **你的工作流程**:
    1.  仔细阅读“最新创作的章节内容”。
    2.  识别其中出现的任何**新信息**，例如：
        - **新登场的人物**: 姓名、身份、性格等关键信息。
        - **对现有人物状态的重大更新**: 目标改变、获得新能力、关系发生重大变化等。
        - **新出现的地点、物品或概念**。
        - **新揭示的世界规则或背景故事**。
    3.  将这些新信息或变化，以简洁的Markdown列表形式输出。**只输出新增和变化的部分，不要重复世界观中已有的内容。**

    **输出示例**:
    - **林渊**: 在第九章遇到了影璃，确认了她是命格守护者之一，并得知了“九渊之地”是下一步的目标。
    - **新登场人物 - 影璃**: 命格守护者之一，引导林渊，告诉他“无命”是起点，并指引他前往“九渊之地”。
    - **新地点 - 枯树林**: 命格之塔的残骸所在地。

    现在，请开始提取章节《{section_title}》中的世界观更新点。
    """
}

# --- Web App Setup ---
app = Flask(__name__)
log_queue = queue.Queue()
task_manager = None 
token_tracker = None

def log_message(message, task_id=None):
    """Logs a message to the queue, optionally associating it with a task ID."""
    log_entry = {"timestamp": time.time(), "message": message, "task_id": task_id}
    log_queue.put(log_entry)

# --- Core Logic Classes (V5) ---

class TokenTracker:
    def __init__(self, filepath):
        self.filepath = filepath
        self.lock = threading.Lock()
        self.stats = self._load()

    def _load(self):
        try:
            if os.path.exists(self.filepath):
                with open(self.filepath, 'r', encoding='utf-8') as f:
                    return json.load(f)
        except (IOError, json.JSONDecodeError):
            pass
        return {"total_prompt_tokens": 0, "total_completion_tokens": 0, "daily_stats": {}}

    def _save(self):
        with open(self.filepath, 'w', encoding='utf-8') as f:
            json.dump(self.stats, f, indent=2)

    def update(self, prompt_tokens, completion_tokens):
        with self.lock:
            today = datetime.utcnow().strftime('%Y-%m-%d')
            self.stats['total_prompt_tokens'] += prompt_tokens
            self.stats['total_completion_tokens'] += completion_tokens
            
            if today not in self.stats['daily_stats']:
                self.stats['daily_stats'][today] = {"prompt_tokens": 0, "completion_tokens": 0}
            
            self.stats['daily_stats'][today]['prompt_tokens'] += prompt_tokens
            self.stats['daily_stats'][today]['completion_tokens'] += completion_tokens
            
            self._save()
    
    def get_stats(self):
        with self.lock:
            return self.stats

class ContentGenerator:
    """Handles all interactions with the LLM API."""
    def __init__(self, api_url, api_headers, model_id, tracker):
        self.api_url = api_url
        self.api_headers = api_headers
        self.model_id = model_id
        self.tracker = tracker

    def _call_llm(self, prompt, max_tokens=8192, temperature=0.75, task_id=None):
        log_message(f"  > Calling LLM API...", task_id)
        payload = { "model": self.model_id, "messages": [{"role": "system", "content": "You are a creative and intelligent assistant."}, {"role": "user", "content": prompt}], "max_tokens": max_tokens, "temperature": temperature }
        try:
            response = requests.post(self.api_url, headers=self.api_headers, data=json.dumps(payload), timeout=600)
            response.raise_for_status()
            data = response.json()
            content = data['choices'][0]['message']['content']
            
            # Track token usage
            usage = data.get('usage', {})
            prompt_tokens = usage.get('prompt_tokens', 0)
            completion_tokens = usage.get('completion_tokens', 0)
            if self.tracker and (prompt_tokens > 0 or completion_tokens > 0):
                self.tracker.update(prompt_tokens, completion_tokens)
            
            log_message(f"  > LLM response received. Tokens used: P={prompt_tokens}, C={completion_tokens}", task_id)
            return content
        except requests.exceptions.RequestException as e:
            log_message(f"  [ERROR] API request failed: {e}", task_id)
            return f"API_REQUEST_ERROR: {e}"
        except (KeyError, IndexError) as e:
            log_message(f"  [ERROR] Failed to parse API response: {e}", task_id)
            return f"API_RESPONSE_ERROR: {e}"

    def generate_outline(self, topic, outline_type, structure, task_id=None):
        log_message(f"--- Starting AI Outline Generation for '{topic}' ---", task_id)
        prompt_template = OUTLINE_PROMPT_TEMPLATES[outline_type]['default']
        if outline_type == 'novel':
            prompt = prompt_template.format(topic=topic, chapters=structure.get('chapters', 5))
        else: # course
            prompt = prompt_template.format(topic=topic, chapters=structure.get('chapters', 5), sections=structure.get('sections', 4))
        outline = self._call_llm(prompt, task_id=task_id)
        log_message("--- Outline generated. ---", task_id)
        return outline

    def generate_initial_worldview(self, topic, outline, task_id=None):
        log_message(f"--- Generating Initial Worldview for '{topic}' ---", task_id)
        prompt = CONTENT_PROMPT_TEMPLATES['initial_worldview'].format(topic=topic, outline=outline)
        worldview = self._call_llm(prompt, task_id=task_id)
        log_message("--- Initial Worldview generated. ---", task_id)
        return worldview

    def generate_chapter_script(self, worldview, previous_chapter_summary, section_title, task_id=None):
        log_message(f"  > Generating Mind Library script for '{section_title}'...", task_id)
        prompt = CONTENT_PROMPT_TEMPLATES['generate_chapter_script'].format(
            worldview=worldview,
            previous_chapter_summary=previous_chapter_summary,
            section_title=section_title
        )
        script_str = self._call_llm(prompt, max_tokens=2048, task_id=task_id)
        
        json_match = re.search(r'\{.*\}', script_str, re.DOTALL)
        if not json_match:
            log_message("  [ERROR] Failed to find valid JSON object in chapter script response.", task_id)
            return None
        
        try:
            script_obj = json.loads(json_match.group(0))
            log_message(f"  > Script for '{section_title}' generated successfully.", task_id)
            return script_obj
        except json.JSONDecodeError as e:
            log_message(f"  [ERROR] Failed to parse chapter script JSON: {e}", task_id)
            return None

    def extract_worldview_updates(self, existing_worldview, new_content, section_title, task_id=None):
        log_message("  > Extracting worldview updates from new content...", task_id)
        context_worldview = (existing_worldview if len(existing_worldview) < 4000 else '...' + existing_worldview[-4000:])
        prompt = CONTENT_PROMPT_TEMPLATES['extract_worldview_updates'].format(
            existing_worldview=context_worldview, 
            new_content=new_content,
            section_title=section_title
        )
        updates = self._call_llm(prompt, max_tokens=2048, task_id=task_id)
        log_message("  > Worldview updates extracted successfully.", task_id)
        return updates

    def _generate_novel_section(self, task, chapter_plan_obj):
        task_id = task['id']
        log_message(f"--- Generating novel chapter: '{task['section_title']}' ---", task_id)

        if not chapter_plan_obj:
            log_message(f"  [ERROR] Chapter plan was not provided for '{task['section_title']}'.", task_id)
            return f"# {task['topic']}\n\n## {task['section_title']}\n\n*Content generation failed: Chapter plan not provided.*"
        
        chapter_plan = json.dumps(chapter_plan_obj, ensure_ascii=False, indent=2)

        previous_chapter_context = "这是小说的第一章。"
        if task.get('previous_content'):
            log_message("  > Summarizing previous chapter for context...", task_id)
            summary_prompt = CONTENT_PROMPT_TEMPLATES['summarizer'].format(text_to_summarize=task['previous_content'])
            summary = self._call_llm(summary_prompt, max_tokens=500, task_id=task_id)
            if "API_" not in summary:
                previous_chapter_context = summary

        worldview_context = "无"
        worldview_path = os.path.join(task['base_dir'], WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
        if os.path.exists(worldview_path):
            log_message("  > Loading worldview context...", task_id)
            with open(worldview_path, 'r', encoding='utf-8') as f:
                worldview_context = f.read()

        log_message(f"  > Generating content for '{task['section_title']}' based on Mind Library plan...", task_id)
        prompt = CONTENT_PROMPT_TEMPLATES['novel'].format(
            topic=task['topic'],
            section_title=task['section_title'],
            word_count=task.get('word_count', 2500),
            previous_chapter_context=previous_chapter_context,
            worldview_context=worldview_context,
            chapter_plan=chapter_plan
        )
        content = self._call_llm(prompt, task_id=task_id)
        if "API_" in content:
            return f"# {task['topic']}\n\n## {task['section_title']}\n\n*Content generation failed.*\n\nError: {content}"
        
        formatted_content = f"# {task['topic']}\n\n## {task['section_title']}\n\n{content}"

        log_message(f"  > Content generation successful for '{task['section_title']}'!", task_id)
        time.sleep(1)
        return formatted_content

    def generate_content_for_section(self, task, chapter_plan_obj=None):
        if task['outline_type'] == 'novel':
            return self._generate_novel_section(task, chapter_plan_obj)
        
        task_id = task['id']
        log_message(f"  > Generating content for '{task['section_title']}'...", task_id)
        prompt_template = CONTENT_PROMPT_TEMPLATES.get(task['outline_type'])
        prompt = prompt_template.format(topic=task['topic'], section_title=task['section_title'])
        content = self._call_llm(prompt, task_id=task_id)
        if "API_" in content:
            return f"# {task['section_title']}\n\n*Content generation failed.*\n\nError: {content}"
        log_message(f"  > Content generation successful!", task_id)
        time.sleep(1)
        return content


class Task:
    """A class to represent a single generation task."""
    def __init__(self, topic, outline, outline_type, word_count=None):
        self.id = str(uuid.uuid4())
        self.topic = topic
        self.outline = outline
        self.outline_type = outline_type
        self.word_count = word_count
        self.status = "queued"
        self.is_cancellation_requested = False
        self.base_dir = os.path.join(WORKS_LIBRARY_PATH, self.outline_type, self.sanitize_filename(self.topic))
        self.sub_tasks = self._create_sub_tasks()
        if not self.sub_tasks: raise ValueError("Failed to parse outline or no sub-tasks were created.")

    @staticmethod
    def sanitize_filename(name, max_length=100):
        name = re.sub(r'^\d+\.\s*', '', name)
        name = name.replace('**', '')
        sanitized_name = re.sub(r'[\\/*?:"<>|]', '_', name)
        return sanitized_name[:max_length]

    def _create_sub_tasks(self):
        sub_tasks = []
        current_chapter_dir = ""
        
        if self.outline_type == 'novel':
            chapter_counter = 1
            for line in self.outline.strip().split('\n'):
                line = line.strip()
                if not line or line.lower().startswith("书名："):
                    continue
                if line.startswith('###'):
                    chapter_title = line.replace('###', '').strip()
                    file_name = f"第{chapter_counter:03d}章-{self.sanitize_filename(chapter_title)}.md"
                    sub_tasks.append({
                        "id": self.id, "topic": self.topic, "outline_type": self.outline_type,
                        "section_title": chapter_title,
                        "file_path": os.path.join(self.base_dir, file_name),
                        "base_dir": self.base_dir, "word_count": self.word_count,
                        "status": "pending"
                    })
                    chapter_counter += 1
        else: # Course logic
            for line in self.outline.strip().split('\n'):
                line = line.strip()
                if not line or line.lower().startswith("标题："):
                    continue
                if line.startswith('###'):
                    chapter_title = self.sanitize_filename(line.replace('###', '').strip())
                    current_chapter_dir = os.path.join(self.base_dir, chapter_title)
                elif re.match(r'^\d+\.', line) and current_chapter_dir:
                    section_title = line.strip()
                    file_name = self.sanitize_filename(section_title) + ".md"
                    sub_tasks.append({
                        "id": self.id, "topic": self.topic, "outline_type": self.outline_type,
                        "section_title": section_title,
                        "file_path": os.path.join(current_chapter_dir, file_name),
                        "base_dir": self.base_dir, "word_count": self.word_count,
                        "status": "pending"
                    })
        return sub_tasks

    def to_dict(self):
        return {
            "id": self.id,
            "topic": self.topic,
            "outline": self.outline,
            "outline_type": self.outline_type,
            "word_count": self.word_count,
            "status": self.status,
            "is_cancellation_requested": self.is_cancellation_requested,
            "base_dir": self.base_dir,
            "sub_tasks": self.sub_tasks
        }

    @classmethod
    def from_dict(cls, data):
        task = cls(data['topic'], data['outline'], data['outline_type'], data.get('word_count'))
        task.id = data['id']
        task.status = data['status']
        task.is_cancellation_requested = data['is_cancellation_requested']
        task.base_dir = data['base_dir']
        task.sub_tasks = data['sub_tasks']
        return task

class TaskManager:
    """Manages the queue and execution of generation tasks."""
    def __init__(self, generator):
        self.task_queue = queue.Queue()
        self.generator = generator
        self.is_running = False
        self.thread = None
        self.current_task = None
        self.lock = threading.Lock()
        self.tasks = {}
        self.tasks_state_file = os.path.join(WORKS_LIBRARY_PATH, TASKS_STATE_FILE)
        self.new_tasks_state_file = os.path.join(WORKS_LIBRARY_PATH, NEW_TASKS_STATE_FILE) # New constant used
        self.old_tasks_state_file = os.path.join(WORKS_LIBRARY_PATH, OLD_TASKS_STATE_FILE) # New constant used
        self._load_tasks_state()

    def _save_tasks_state(self):
        with self.lock:
            try:
                with open(self.tasks_state_file, 'w', encoding='utf-8') as f:
                    json.dump({tid: task.to_dict() for tid, task in self.tasks.items()}, f, ensure_ascii=False, indent=4)
            except IOError as e:
                log_message(f"[ERROR] Could not save tasks state: {e}")

    def _load_tasks_state(self):
        if not os.path.exists(self.tasks_state_file):
            return
        try:
            with open(self.tasks_state_file, 'r', encoding='utf-8') as f:
                tasks_data = json.load(f)
            for tid, tdata in tasks_data.items():
                task = Task.from_dict(tdata)
                self.tasks[tid] = task
                if task.status in ["queued", "running"]:
                    task.status = "queued" # Re-queue previously running tasks
                    self.task_queue.put(task)
                    print(f"Restored and re-queued task '{task.topic}'.", task.id)
                    log_message(f"Restored and re-queued task '{task.topic}'.", task.id)
        except (IOError, json.JSONDecodeError) as e:
            log_message(f"[ERROR] Could not load tasks state: {e}")

    def worker(self):
        while self.is_running:
            try:
                task = self.task_queue.get(timeout=1)
                with self.lock:
                    self.current_task = task
                    self.current_task.status = "running"
                self._save_tasks_state()
                log_message(f"--- Starting task: {task.topic} (ID: {task.id}) ---", task.id)
                
                mind_library = []
                mind_library_path = None
                if task.outline_type == 'novel':
                    mind_library_dir = os.path.join(task.base_dir, MIND_LIBRARY_DIR_NAME)
                    if not os.path.exists(mind_library_dir): os.makedirs(mind_library_dir)
                    mind_library_path = os.path.join(mind_library_dir, MIND_LIBRARY_FILE_NAME)
                    if os.path.exists(mind_library_path):
                        with open(mind_library_path, 'r', encoding='utf-8') as f:
                            mind_library = json.load(f)
                    # 增加每一篇小说的task，保存好，才能跟task_state.json 联动重试。。
                    j = os.path.join(mind_library_dir, TASK_LIBRARY_FILE_NAME)
                    with open(j, 'w', encoding='utf-8') as f:
                        json.dump({"taskId": task.id}, f, ensure_ascii=False,
                                  indent=4)

                previous_content = None
                previous_summary = "这是小说的第一章。"
                task_failed = False
                for index, sub_task in enumerate(task.sub_tasks):
                    if task.is_cancellation_requested:
                        task.status = "cancelled"
                        log_message(f"Task '{task.topic}' cancelled.", task.id)
                        break
                    
                    if sub_task.get('status') == 'completed':
                        log_message(f"  - [Skipping] Sub-task already completed: {sub_task['section_title']}", task.id)
                        # Ensure previous_content is updated for the next chapter
                        if task.outline_type == 'novel' and os.path.exists(sub_task['file_path']):
                             with open(sub_task['file_path'], 'r', encoding='utf-8') as f: 
                                previous_content = f.read()
                        continue
                    
                    if not os.path.exists(os.path.dirname(sub_task['file_path'])): 
                        os.makedirs(os.path.dirname(sub_task['file_path']))
                    
                    if os.path.exists(sub_task['file_path']):
                        log_message(f"  - [Skipping] File exists: {sub_task['file_path']}", task.id)
                        sub_task['status'] = 'completed'
                        self._save_tasks_state()
                        if task.outline_type == 'novel':
                            with open(sub_task['file_path'], 'r', encoding='utf-8') as f: 
                                previous_content = f.read()
                        continue
                    
                    sub_task['previous_content'] = previous_content
                    
                    chapter_plan_obj = None
                    max_retries = 1
                    for attempt in range(max_retries + 1):
                        try:
                            if task.outline_type == 'novel':
                                # Just-in-time script generation
                                if index >= len(mind_library):
                                    worldview_path = os.path.join(task.base_dir, WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
                                    worldview_context = ""
                                    if os.path.exists(worldview_path):
                                        with open(worldview_path, 'r', encoding='utf-8') as f:
                                            worldview_context = f.read()
                                    
                                    if previous_content:
                                        summary = self.generator._call_llm(CONTENT_PROMPT_TEMPLATES['summarizer'].format(text_to_summarize=previous_content), max_tokens=500, task_id=task.id)
                                        if "API_" not in summary: previous_summary = summary

                                    new_script = self.generator.generate_chapter_script(worldview_context, previous_summary, sub_task['section_title'], task.id)
                                    if new_script:
                                        mind_library.append(new_script)
                                        with open(mind_library_path, 'w', encoding='utf-8') as f:
                                            json.dump(mind_library, f, ensure_ascii=False, indent=2)
                                        chapter_plan_obj = new_script
                                    else:
                                        log_message(f"  [ERROR] Failed to generate script for '{sub_task['section_title']}'. Skipping chapter.", task.id)
                                        raise Exception(f"Failed to generate script for '{sub_task['section_title']}'")
                                else:
                                    chapter_plan_obj = mind_library[index]
                            
                            content = self.generator.generate_content_for_section(sub_task, chapter_plan_obj)
                            if "generation failed" in content:
                                raise Exception("LLM content generation failed.")

                            with open(sub_task['file_path'], 'w', encoding='utf-8') as f: 
                                f.write(content)
                            log_message(f"    - [Success] Saved file: {sub_task['file_path']}", task.id)
                            previous_content = content
                            
                            if task.outline_type == 'novel': # No retry for these secondary tasks
                                worldview_path = os.path.join(task.base_dir, WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
                                if os.path.exists(worldview_path):
                                    with open(worldview_path, 'r', encoding='utf-8') as f: 
                                        existing_worldview = f.read()
                                    
                                    updates = self.generator.extract_worldview_updates(existing_worldview, content, sub_task['section_title'], task.id)
                                    
                                    if "API_" not in updates and updates.strip():
                                        with open(worldview_path, 'a', encoding='utf-8') as f:
                                            f.write(f"\n\n---\n\n### 《{sub_task['section_title']}》章节更新\n\n{updates}")
                                        log_message(f"    - [Success] Appended updates to Worldview.", task.id)
                                
                                if mind_library and mind_library_path and index < len(mind_library):
                                    mind_library[index]["status"] = "completed"
                                    with open(mind_library_path, 'w', encoding='utf-8') as f:
                                        json.dump(mind_library, f, ensure_ascii=False, indent=2)
                                    log_message(f"    - [Success] Updated Mind Library status for '{sub_task['section_title']}'", task.id)
                            break # Success, exit retry loop
                        except Exception as e:
                            log_message(f"    - [Attempt {attempt + 1}/{max_retries + 1}] Failed for '{sub_task['section_title']}'. Error: {e}", task.id)
                            if attempt < max_retries:
                                time.sleep(3) # Wait before retrying
                            else:
                                log_message(f"    - [Failed] Max retries reached for '{sub_task['section_title']}'. Task will be marked as failed.", task.id)
                                task.status = "failed"
                                task_failed = True
                                break # Exit retry loop

                    if task_failed:
                        break # Exit sub-task loop

                if task.status not in ["cancelled", "failed"]:
                    task.status = "completed"
                log_message(f"--- Task '{task.topic}' {task.status} ---", task.id)
                
                with self.lock:
                    self.tasks[task.id] = task
                    self.current_task = None
                self._save_tasks_state()
                self.task_queue.task_done()
            except queue.Empty: 
                continue
        log_message("Worker thread has stopped.")

    def backup_file(self, file_path):
        """
        备份文件，添加时间戳作为后缀

        :param file_path: 原始文件路径
        :return: 新文件路径，如果备份失败返回None
        """
        if not os.path.exists(file_path):
            return None

        # 获取文件目录和文件名
        dir_name = os.path.dirname(file_path)
        file_name = os.path.basename(file_path)

        # 分离文件名和扩展名
        name, ext = os.path.splitext(file_name)

        # 添加时间戳
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        new_file_name = f"{name}.{timestamp}{ext}"
        new_file_path = os.path.join(dir_name, new_file_name)

        try:
            # 重命名文件
            os.rename(file_path, new_file_path)
            return new_file_path
        except OSError as e:
            print(f"备份文件失败: {e}")
            return None

    def retry_task(self, task_id, file_name):
        with self.lock:
            task = self.tasks.get(task_id)
            if not task:
                return False, "Task not found."
            # if task.status != "failed": TODO
            #     return False, "Only failed tasks can be retried."

            task.status = "queued"
            task.is_cancellation_requested = False
            # 主要改第几章了
            # 搜索subTasks 下的section_title 包含了file_name
            # 搜索到了才需要重试
            retry = False
            for subtask in task.sub_tasks:
                subtitle = subtask["section_title"]
                subtitle = subtitle.replace('*', '')
                print(subtitle)
                if subtitle in file_name:
                   subtask["status"] ="queued"
                   retry = True
                   # 似乎要删除文件才行，所以改为做备份文件吧
                   self.backup_file(subtask["file_path"])
            if not retry:
                return False, "Task dont find"
            self.task_queue.put(task)
            log_message(f"Task '{task.topic}' has been retried.", task.id)
            return True, "Task has been retried successfully."

    def start(self):
        if not self.is_running:
            self.is_running = True
            self.thread = threading.Thread(target=self.worker)
            self.thread.daemon = True
            self.thread.start()
            log_message("Task manager started.")

    def add_task(self, task):
        with self.lock: self.tasks[task.id] = task
        self.task_queue.put(task)
        self._save_tasks_state()
        log_message(f"Task '{task.topic}' added to queue.", task.id)

    def cancel_task(self, task_id):
        with self.lock:
            task = self.tasks.get(task_id)
            if not task: return False, "Task not found."
            if task.status == "running":
                task.is_cancellation_requested = True
                return True, "Cancellation requested."
            if task.status == "queued":
                task.status = "cancelled"
                new_q = queue.Queue()
                while not self.task_queue.empty():
                    item = self.task_queue.get()
                    if item.id != task_id: new_q.put(item)
                self.task_queue = new_q
                log_message(f"Task '{task.topic}' removed from queue.", task.id)
                self._save_tasks_state()
                return True, "Task removed from queue."
            return False, f"Task is already {task.status}."

    def get_tasks_status(self):
        with self.lock:
            status_list = []
            if self.current_task: status_list.append({"id": self.current_task.id, "topic": self.current_task.topic, "status": self.current_task.status})
            for task in list(self.task_queue.queue): status_list.append({ "id": task.id, "topic": task.topic, "status": task.status })
            for task in self.tasks.values():
                if task.status in ['completed', 'cancelled', 'failed'] and task.id not in [t['id'] for t in status_list]: status_list.append({ "id": task.id, "topic": task.topic, "status": task.status })
            return status_list

# --- HTML Template (V5 - Added Token Stats) ---
with open('index.html', 'r', encoding='utf-8') as file:
    HTML_TEMPLATE = file.read()

# --- Helper function for file conversion ---
def markdown_to_plain_text(md_content):
    """Converts Markdown content to plain text."""
    content_sans_titles = re.sub(r'^# .*\n', '', md_content)
    content_sans_titles = re.sub(r'^## .*\n', '', content_sans_titles)
    html = markdown(content_sans_titles)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text()

def markdown_to_docx(md_content):
    """Converts Markdown content to a DOCX file in memory."""
    document = Document()
    html = markdown(md_content)
    soup = BeautifulSoup(html, "html.parser")
    for element in soup.contents:
        if element.name == 'h1': document.add_heading(element.text, level=1)
        elif element.name == 'h2': document.add_heading(element.text, level=2)
        elif element.name == 'h3': document.add_heading(element.text, level=3)
        elif element.name == 'p': document.add_paragraph(element.text)
        elif element.name in ['ul', 'ol']:
            for li in element.find_all('li'): document.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'pre': document.add_paragraph(element.text)
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream

# --- Flask Routes (V5 - Upgraded) ---

@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/generate-outline', methods=['POST'])
def generate_outline_api():
    data = request.json
    generator = ContentGenerator(TARGET_API_URL, API_HEADERS, SERVED_MODEL_IDENTIFIER, token_tracker)
    structure = {"chapters": int(data.get('chapters', 10)), "sections": int(data.get('sections', 4))}
    outline = generator.generate_outline(data['topic'], data['outline_type'], structure)
    if outline and "API_" not in outline:
        return jsonify({"outline": outline})
    return jsonify({"error": "Failed to generate outline"}), 500

@app.route('/start-task', methods=['POST'])
def start_task_api():
    data = request.json
    outline = data.get('outline')
    original_topic = data.get('original_topic')
    outline_type = data.get('outline_type')
    if not all([outline, original_topic, outline_type]):
        return jsonify({"status": "error", "error": "Missing data"}), 400
    
    first_line = outline.split('\n', 1)[0]
    title_match = re.search(r'^(?:标题|书名)：\s*(.+)', first_line.strip(), re.IGNORECASE)
    topic = title_match.group(1).strip() if title_match else Task.sanitize_filename(original_topic, max_length=50)

    if outline_type == 'novel':
        generator = ContentGenerator(TARGET_API_URL, API_HEADERS, SERVED_MODEL_IDENTIFIER, token_tracker)
        log_message(f"Creating Worldview for {topic}...")
        initial_worldview = generator.generate_initial_worldview(topic, outline)
        if "API_" in initial_worldview:
            return jsonify({"status": "error", "error": "Failed to generate initial worldview"}), 500
        
        project_base_dir = os.path.join(WORKS_LIBRARY_PATH, 'novel', Task.sanitize_filename(topic))
        worldview_dir = os.path.join(project_base_dir, WORLDVIEW_DIR_NAME)
        if not os.path.exists(worldview_dir): os.makedirs(worldview_dir)
        with open(os.path.join(worldview_dir, WORLDVIEW_FILE_NAME), 'w', encoding='utf-8') as f:
            f.write(initial_worldview)
        
        # Mind library is now generated just-in-time, so we don't create it here.
        # We just create the directory.
        mind_library_dir = os.path.join(project_base_dir, MIND_LIBRARY_DIR_NAME)
        if not os.path.exists(mind_library_dir): os.makedirs(mind_library_dir)


    try:
        task = Task(topic, outline, outline_type, data.get('word_count'))
        task_manager.add_task(task)
        return jsonify({"status": "success", "message": "Task added to queue", "task_id": task.id})
    except ValueError as e:
        return jsonify({"status": "error", "error": str(e)}), 400

@app.route('/status')
def status_api():
    logs = []
    while not log_queue.empty():
        logs.append(log_queue.get_nowait())
    return jsonify(logs)

@app.route('/tasks')
def get_tasks_api():
    return jsonify(task_manager.get_tasks_status())

@app.route('/tasks/cancel/<task_id>', methods=['POST'])
def cancel_task_api(task_id):
    success, message = task_manager.cancel_task(task_id)
    if success:
        return jsonify({"message": message})
    return jsonify({"error": message}), 400

@app.route('/library')
def get_library_api():
    library = {"course": [], "novel": []}
    if not os.path.exists(WORKS_LIBRARY_PATH):
        os.makedirs(os.path.join(WORKS_LIBRARY_PATH, "course"))
        os.makedirs(os.path.join(WORKS_LIBRARY_PATH, "novel"))
    for type in ["course", "novel"]:
        type_path = os.path.join(WORKS_LIBRARY_PATH, type)
        if os.path.exists(type_path):
            library[type] = sorted([d for d in os.listdir(type_path) if os.path.isdir(os.path.join(type_path, d))])
    return jsonify(library)

@app.route('/library/<type>/<project_name>')
def get_project_details_api(type, project_name):
    project_path = os.path.join(WORKS_LIBRARY_PATH, type, project_name)
    if not os.path.exists(project_path): return jsonify({"error": "Project not found"}), 404
    
    if type == 'novel':
        # 获取_mind_library 中的task.json 中的taskId 一同返回
        taskFile = os.path.join(project_path, MIND_LIBRARY_DIR_NAME, TASK_LIBRARY_FILE_NAME)
        with open(taskFile, 'r', encoding='utf-8') as f:
            taskId = json.load(f)
        files = sorted([f for f in os.listdir(project_path) if f.endswith('.md')])
        return jsonify({"name": project_name, "tree": files, "taskId": taskId})
    else: # Course
        tree = {}
        for chapter in sorted([d for d in os.listdir(project_path) if os.path.isdir(os.path.join(project_path, d)) and d != WORLDVIEW_DIR_NAME]):
            chapter_path = os.path.join(project_path, chapter)
            files = sorted([f for f in os.listdir(chapter_path) if f.endswith('.md')])
            tree[chapter] = files
        return jsonify({"name": project_name, "tree": tree})

@app.route('/library/content/<type>/<project_name>/<path:filepath>')
def get_file_content_api(type, project_name, filepath):
    full_path = os.path.join(WORKS_LIBRARY_PATH, type, project_name, filepath)
    if not os.path.exists(full_path): return jsonify({"error": "File not found"}), 404
    with open(full_path, 'r', encoding='utf-8') as f: content = f.read()
    return jsonify({"content": content})

@app.route('/download/zip/<type>/<project_name>')
def download_project_zip_api(type, project_name):
    project_path = os.path.join(WORKS_LIBRARY_PATH, type, project_name)
    if not os.path.exists(project_path): return jsonify({"error": "Project not found"}), 404
    memory_file = BytesIO()
    with shutil.ZipFile(memory_file, 'w', shutil.ZIP_DEFLATED) as zf:
        for root, _, files in os.walk(project_path):
            if WORLDVIEW_DIR_NAME in root or MIND_LIBRARY_DIR_NAME in root: continue
            for file in files:
                zf.write(os.path.join(root, file), os.path.relpath(os.path.join(root, file), os.path.join(project_path, '..')))
    memory_file.seek(0)
    return send_file(memory_file, download_name=f'{project_name}.zip', as_attachment=True)

@app.route('/api/worldview/<project_name>')
def worldview_api(project_name):
    worldview_file = os.path.join(WORKS_LIBRARY_PATH, 'novel', project_name, WORLDVIEW_DIR_NAME, WORLDVIEW_FILE_NAME)
    if os.path.exists(worldview_file):
        with open(worldview_file, 'r', encoding='utf-8') as f: content = f.read()
        return jsonify({"content": content})
    return jsonify({"content": ""})

@app.route('/api/mind-library/<project_name>')
def mind_library_api(project_name):
    mind_library_file = os.path.join(WORKS_LIBRARY_PATH, 'novel', project_name, MIND_LIBRARY_DIR_NAME, MIND_LIBRARY_FILE_NAME)
    if os.path.exists(mind_library_file):
        with open(mind_library_file, 'r', encoding='utf-8') as f: library = json.load(f)
        return jsonify({"library": library})
    return jsonify({"library": []})

@app.route('/api/token-stats')
def token_stats_api():
    return jsonify(token_tracker.get_stats())

@app.route('/api/showcase-sequence/<project_name>/<path:section_name>')
def get_showcase_sequence_api(project_name, section_name):
    project_path = os.path.join(WORKS_LIBRARY_PATH, 'novel', project_name)
    if not os.path.exists(project_path): return jsonify({"error": "Project not found"}), 404
    
    all_chapters = sorted([f for f in os.listdir(project_path) if f.endswith('.md')])
    try:
        start_chapter_index = all_chapters.index(section_name)
    except ValueError:
        return jsonify({"error": "Starting chapter not found"}), 404
        
    sections_to_play = []
    for chapter_idx in range(start_chapter_index, len(all_chapters)):
        current_chapter_name = all_chapters[chapter_idx]
        with open(os.path.join(project_path, current_chapter_name), 'r', encoding='utf-8') as f: 
            content = f.read()
        sections_to_play.append({"name": current_chapter_name, "content": content})
    return jsonify({"sections": sections_to_play})

@app.route('/tasks/retry/<task_id>/<file_name>', methods=['POST'])
def retry_task_api(task_id, file_name):
    # filename 主要是为了找到第几章，retry_task
    success, message = task_manager.retry_task(task_id, file_name)
    if success:
        return jsonify({"message": message})
    return jsonify({"error": message}), 400

# --- Main Execution ---
if __name__ == '__main__':
    for dir_path in [WORKS_LIBRARY_PATH, os.path.join(WORKS_LIBRARY_PATH, "course"), os.path.join(WORKS_LIBRARY_PATH, "novel")]:
        if not os.path.exists(dir_path): os.makedirs(dir_path)
    
    token_tracker = TokenTracker(TOKEN_STATS_FILE)
    generator = ContentGenerator(TARGET_API_URL, API_HEADERS, SERVED_MODEL_IDENTIFIER, token_tracker)
    task_manager = TaskManager(generator)
    task_manager.start()
    
    print("="*50)
    print("AI Content Factory v7.0 (Web UI) Started!")
    print(f"Please open your browser to http://127.0.0.1:48000")
    print("="*50)
    app.run(host='0.0.0.0', port=48000, debug=False)
